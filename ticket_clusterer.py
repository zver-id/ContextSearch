from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sentence_transformers import SentenceTransformer
from sklearn.cluster import KMeans, DBSCAN
from sklearn.metrics import silhouette_score
from sklearn.neighbors import NearestNeighbors
from sklearn.preprocessing import StandardScaler
import numpy as np
from typing import Dict, List, Union, Any, Tuple, Optional
import matplotlib.pyplot as plt
import re
from collections import Counter
import nltk
from nltk.corpus import stopwords, names
import logging
import spacy


# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TicketClusterer:
    """
    Кластеризатор текстов с поддержкой коротких и длинных текстов
    """

    def __init__(self, model_name: str = 'sentence-transformers/paraphrase-multilingual-mpnet-base-v2',
                 max_tokens: int = 512):
        """
        Инициализация
        :param model_name: название модели SentenceTransformer
        :param max_tokens: максимальное количество токенов для коротких текстов
        """
        self.model = SentenceTransformer(model_name)
        self.max_tokens = max_tokens
        self.embeddings = None
        self.document_ids = None
        self.documents = None
        self.cluster_labels = None
        self.cluster_model = None
        self.chunk_info = None
        self.processing_mode = None  # 'short' или 'long'
        self.nlp = nlp = spacy.load("ru_core_news_sm")

        # Загрузка стоп-слов

        nltk.download('stopwords', quiet=True)
        nltk.download('names', quiet=True)
        #names_set = set(names.words('male.txt') + names.words('female.txt'))
        names_set = set()
        stop_words = set(stopwords.words('russian') + stopwords.words('english'))
        self.stop_words = stop_words.union(names_set)

#region Методы для коротких текстов (до 512 токенов)

    def cluster_short_texts(self, id_text_dict: Dict[Any, str],
                            n_clusters: int = None,
                            method: str = "kmeans",
                            **kwargs) -> Dict[str, Any]:
        """
        Кластеризация коротких текстов (до 512 токенов)
        :param id_text_dict: словарь {id: текст}
        :param n_clusters: количество кластеров (если None, определяется автоматически)
        :param method: метод кластеризации ("kmeans", "dbscan")
        :param kwargs: дополнительные параметры для методов кластеризации
        :return: результаты кластеризации
        """
        self.processing_mode = 'short'
        logger.info("Режим: кластеризация коротких текстов")
        self._prepare_basic_data(id_text_dict)
        self._validate_text_lengths()
        self._create_simple_embeddings(normalize=True)

        if method == "kmeans":
            return self._kmeans_cluster_short(n_clusters, **kwargs)
        elif method == "dbscan":
            return self._dbscan_cluster_short(**kwargs)
        else:
            raise ValueError(f"Неизвестный метод кластеризации: {method}")

    def _validate_text_lengths(self) -> None:
        """
        Проверяет, что тексты достаточно короткие
        """
        for doc_id, text in zip(self.document_ids, self.documents):
            token_count = len(text.split())
            if token_count > self.max_tokens * 1.5:  # Небольшой запас
                logger.warning(f"Текст {doc_id} содержит {token_count} токенов. "
                               f"Рекомендуется использовать cluster_long_texts()")

    def _create_simple_embeddings(self, normalize: bool = True) -> np.ndarray:
        """
        Создание эмбеддингов для коротких текстов
        :param normalize: нормализация данных
        :return: эмбендинги коротких текстов
        """

        logger.info("Создание эмбеддингов для коротких текстов...")

        self.embeddings = self.model.encode(
            self.documents,
            show_progress_bar=True,
            normalize_embeddings=normalize
        )
        return self.embeddings

    def _kmeans_cluster_short(self, n_clusters: int = None,
                              max_clusters: int = 10,
                              random_state: int = 42) -> Dict[str, Any]:
        """
        K-means для коротких текстов
        :param n_clusters: количество кластеров
        :param max_clusters:
        :param random_state:
        :return:
        """

        if n_clusters is None:
            n_clusters = self._find_optimal_clusters(max_clusters)

        self.cluster_model = KMeans(n_clusters=n_clusters, random_state=random_state)
        self.cluster_labels = self.cluster_model.fit_predict(self.embeddings)

        return self._format_results()

    def _dbscan_cluster_short(self, eps: float = None,
                              min_samples: int = 2) -> Dict[str, Any]:
        """
        DBSCAN для коротких текстов
        :param eps: радиус поиска
        :param min_samples: минимальное количество элементов в кластере
        :return: кластеризованные данные.
        """
        scaler = StandardScaler()
        scaled_embeddings = scaler.fit_transform(self.embeddings)

        if eps is None:
            eps = self._find_optimal_eps(self.embeddings, min_samples)

        self.cluster_model = DBSCAN(eps=eps, min_samples=min_samples)
        self.cluster_labels = self.cluster_model.fit_predict(scaled_embeddings)

        return self._format_results()
# endregion
# region методы для длинных текстов (более 512 токенов)

    def cluster_long_texts(self, id_text_dict: Dict[Any, str],
                           n_clusters: int = None,
                           method: str = "kmeans",
                           chunk_strategy: str = "sliding",
                           aggregation: str = "mean",
                           use_hierarchical: bool = True,
                           **kwargs) -> Dict[str, Any]:
        """
        Кластеризация длинных текстов (более 512 токенов)
        :param id_text_dict: словарь {id: текст}
        :param n_clusters: количество кластеров
        :param method: метод кластеризации ("kmeans", "dbscan")
        :param chunk_strategy: стратегия разбиения ("sliding", "sentences", "paragraphs")
        :param aggregation: метод агрегации ("mean", "max", "weighted")
        :param use_hierarchical: использовать иерархические эмбеддинги
        :param kwargs: дополнительные параметры
        :return: результаты кластеризации
        """
        self.processing_mode = 'long'
        logger.info("Режим: кластеризация длинных текстов")
        self._prepare_basic_data(id_text_dict)

        # Создание эмбеддингов для длинных текстов
        if use_hierarchical:
            self._create_hierarchical_embeddings()
        else:
            self._create_long_text_embeddings(chunk_strategy, aggregation)

        # Кластеризация
        if method == "kmeans":
            return self._kmeans_cluster_long(n_clusters, **kwargs)
        elif method == "dbscan":
            return self._dbscan_cluster_long(**kwargs)
        else:
            raise ValueError(f"Неизвестный метод кластеризации: {method}")

    def _create_long_text_embeddings(self, chunk_strategy: str = "sentences",
                                     aggregation: str = "mean") -> np.ndarray:
        """
        Создание эмбеддингов для длинных текстов через разбиение на чанки
        :param chunk_strategy: стратегия разбиения ("sliding", "sentences", "paragraphs")
        :param aggregation: метод агрегации ("mean", "max", "weighted")
        :return: эмбендинги длинных текстов
        """
        logger.info("Разбиение длинных текстов на чанки...")

        all_chunks = []
        self.chunk_info = []

        for doc_id, text in zip(self.document_ids, self.documents):
            chunks = self._split_into_chunks(text, doc_id, chunk_strategy)
            for chunk_text, metadata in chunks:
                all_chunks.append(chunk_text)
                self.chunk_info.append(metadata)

        logger.info(f"Создано {len(all_chunks)} чанков из {len(self.documents)} документов")

        if not all_chunks:
            logger.warning("Не создано ни одного чанка. Создаем эмбеддинги из исходных текстов.")
            return self._create_simple_embeddings(normalize=True)

        # Эмбеддинги для чанков
        chunk_embeddings = self.model.encode(
            all_chunks,
            show_progress_bar=True,
            normalize_embeddings=True
        )

        if np.isnan(chunk_embeddings).any():
            logger.warning("Обнаружены NaN значения в эмбеддингах чанков. Заменяем нулями.")
            chunk_embeddings = np.nan_to_num(chunk_embeddings)

        logger.info("Агрегация эмбеддингов чанков...")
        doc_embeddings = []

        for doc_id in self.document_ids:
            chunk_indices = [i for i, info in enumerate(self.chunk_info)
                             if info['doc_id'] == doc_id]

            if not chunk_indices:
                logger.warning(f"Для документа {doc_id} не создано чанков. Используем нулевой вектор.")
                doc_emb = np.zeros(chunk_embeddings.shape[1])
            else:
                chunk_embs = chunk_embeddings[chunk_indices]

                valid_chunks = []
                for i, emb in enumerate(chunk_embs):
                    if not np.isnan(emb).any():
                        valid_chunks.append(emb)
                    else:
                        logger.warning(f"Чанк {chunk_indices[i]} документа {doc_id} содержит NaN. Пропускаем.")

                if not valid_chunks:
                    logger.warning(f"Для документа {doc_id} нет валидных чанков. Используем нулевой вектор.")
                    doc_emb = np.zeros(chunk_embeddings.shape[1])
                else:
                    valid_chunks = np.array(valid_chunks)
                    doc_emb = self._aggregate_embeddings(valid_chunks, aggregation,
                                                         chunk_indices, all_chunks)
            doc_embeddings.append(doc_emb)

        self.embeddings = np.array(doc_embeddings)

        if np.isnan(self.embeddings).any():
            logger.warning("Обнаружены NaN значения в финальных эмбеддингах. Заменяем нулями.")
            self.embeddings = np.nan_to_num(self.embeddings)


        norms = np.linalg.norm(self.embeddings, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        self.embeddings = self.embeddings / norms

        return self.embeddings

    def _create_hierarchical_embeddings(self, levels: List[str] = None) -> np.ndarray:
        """
        Создание иерархических эмбеддингов
        :param levels:
        :return:
        """
        if levels is None:
            levels = ["keyword", "chunk", "full"]

        all_embeddings = []

        for level in levels:
            if level == "keyword":
                keyword_texts = []
                for text in self.documents:
                    keywords = self._extract_keywords(text, 15)
                    keyword_text = ' '.join(keywords) if keywords else "пустой текст"
                    keyword_texts.append(keyword_text)

                emb = self.model.encode(keyword_texts, normalize_embeddings=True)
                all_embeddings.append(emb)

            elif level == "chunk":
                emb = self._create_long_text_embeddings()
                all_embeddings.append(emb)

            elif level == "full":
                full_texts = []
                for text in self.documents:
                    words = text.split()[:self.max_tokens * 2]
                    full_text = ' '.join(words) if words else "пустой текст"
                    full_texts.append(full_text)

                emb = self.model.encode(full_texts, normalize_embeddings=True)
                all_embeddings.append(emb)

        # Конкатенация и проверка на NaN
        self.embeddings = np.concatenate(all_embeddings, axis=1)

        if np.isnan(self.embeddings).any():
            logger.warning("Обнаружены NaN в иерархических эмбеддингах. Заменяем нулями.")
            self.embeddings = np.nan_to_num(self.embeddings)

        return self.embeddings

    def _validate_and_clean_embeddings(self) -> None:
        """
        Валидация и очистка эмбеддингов от NaN и бесконечных значений
        """
        if self.embeddings is None:
            raise ValueError("Эмбеддинги не созданы")

        # Проверяем на NaN
        if np.isnan(self.embeddings).any():
            nan_count = np.isnan(self.embeddings).sum()
            logger.warning(f"Обнаружено {nan_count} NaN значений в эмбеддингах. Заменяем нулями.")
            self.embeddings = np.nan_to_num(self.embeddings)

        # Проверяем на бесконечные значения
        if np.isinf(self.embeddings).any():
            inf_count = np.isinf(self.embeddings).sum()
            logger.warning(f"Обнаружено {inf_count} бесконечных значений в эмбеддингах. Заменяем нулями.")
            self.embeddings = np.nan_to_num(self.embeddings, nan=0.0, posinf=0.0, neginf=0.0)

        # Проверяем, что все эмбеддинги не нулевые
        zero_embeddings = np.all(self.embeddings == 0, axis=1)
        if zero_embeddings.any():
            zero_count = zero_embeddings.sum()
            logger.warning(f"Обнаружено {zero_count} нулевых эмбеддингов.")

            # Добавляем небольшой шум к нулевым эмбеддингам
            for i in range(len(self.embeddings)):
                if zero_embeddings[i]:
                    self.embeddings[i] = np.random.normal(0, 0.001, self.embeddings.shape[1])
                    logger.info(f"Добавлен шум к нулевому эмбеддингу документа {i}")





    def _split_into_chunks(self, text: str, doc_id: Any,
                           strategy: str) -> List[Tuple[str, Dict]]:
        """
        Разбивает текст на чанки по выбранной стратегии
        :param text: Текст для разбиения
        :param doc_id: ИД документа
        :param strategy: стратегия разбиения ("sliding", "sentences", "paragraphs")
        :return: документ разбитый на чанки
        """

        chunks = []

        if strategy == "sentences":
            sentences = re.split(r'[.!?]+', text)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 10]

            current_chunk = []
            current_length = 0

            for sentence in sentences:
                sentence_length = len(sentence.split())
                if current_length + sentence_length <= self.max_tokens:
                    current_chunk.append(sentence)
                    current_length += sentence_length
                else:
                    if current_chunk:
                        chunk_text = ' '.join(current_chunk)
                        chunks.append((chunk_text, {
                            'doc_id': doc_id, 'chunk_type': 'sentence_based'
                        }))
                    current_chunk = [sentence]
                    current_length = sentence_length

            if current_chunk:
                chunk_text = ' '.join(current_chunk)
                chunks.append((chunk_text, {
                    'doc_id': doc_id, 'chunk_type': 'sentence_based'
                }))

        elif strategy == "paragraphs":
            paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 0]

            for i, paragraph in enumerate(paragraphs):
                if len(paragraph.split()) <= self.max_tokens:
                    chunks.append((paragraph, {
                        'doc_id': doc_id, 'chunk_type': 'paragraph'
                    }))
                else:
                    words = paragraph.split()
                    for j in range(0, len(words), self.max_tokens):
                        chunk_text = ' '.join(words[j:j + self.max_tokens])
                        chunks.append((chunk_text, {
                            'doc_id': doc_id, 'chunk_type': 'paragraph_chunk'
                        }))

        else:  # sliding window
            words = text.split()
            window_size = self.max_tokens
            overlap = window_size // 4

            for i in range(0, len(words), window_size - overlap):
                chunk_text = ' '.join(words[i:i + window_size])
                if len(chunk_text.strip()) > 0:
                    chunks.append((chunk_text, {
                        'doc_id': doc_id, 'chunk_type': 'sliding_window'
                    }))

        return chunks

    def _aggregate_embeddings(self, chunk_embeddings: np.ndarray,
                              method: str, indices: List[int],
                              all_chunks: List[str]) -> np.ndarray:
        """Агрегирует эмбеддинги чанков"""
        if method == "mean":
            return np.mean(chunk_embeddings, axis=0)
        elif method == "max":
            return np.max(chunk_embeddings, axis=0)
        elif method == "weighted":
            weights = [len(all_chunks[i].split()) for i in indices]
            weights = np.array(weights) / sum(weights)
            return np.average(chunk_embeddings, axis=0, weights=weights)
        else:
            return np.mean(chunk_embeddings, axis=0)

    def _kmeans_cluster_long(self, n_clusters: int = None,
                             max_clusters: int = 10,
                             random_state: int = 42) -> Dict[str, Any]:
        """K-means для длинных текстов"""
        if n_clusters is None:
            n_clusters = self._find_optimal_clusters(max_clusters)

        # Валидация и очистка эмбеддингов
        self._validate_and_clean_embeddings()

        # Проверяем, что у нас достаточно данных для кластеризации
        if len(self.embeddings) < n_clusters:
            logger.warning(
                f"Количество документов ({len(self.embeddings)}) меньше количества кластеров ({n_clusters}). Уменьшаем количество кластеров.")
            n_clusters = max(2, len(self.embeddings) - 1)

        self.cluster_model = KMeans(n_clusters=n_clusters, random_state=random_state)
        self.cluster_labels = self.cluster_model.fit_predict(self.embeddings)

        return self._format_results()

    def _dbscan_cluster_long(self, eps: float = None,
                             min_samples: int = 2) -> Dict[str, Any]:
        """DBSCAN для длинных текстов"""
        # Валидация и очистка эмбеддингов
        self._validate_and_clean_embeddings()

        scaler = StandardScaler()
        scaled_embeddings = scaler.fit_transform(self.embeddings)

        if eps is None:
            eps = self._find_optimal_eps(self.embeddings, min_samples)

        self.cluster_model = DBSCAN(eps=eps, min_samples=min_samples)
        self.cluster_labels = self.cluster_model.fit_predict(scaled_embeddings)

        return self._format_results()

    def cluster_long_texts_with_fallback(self, id_text_dict: Dict[Any, str],
                                         n_clusters: int = None,
                                         method: str = "kmeans",
                                         **kwargs) -> Dict[str, Any]:
        """
        Кластеризация длинных текстов с fallback на короткие тексты при ошибках
        """
        try:
            return self.cluster_long_texts(id_text_dict, n_clusters, method, **kwargs)
        except Exception as e:
            logger.error(f"Ошибка при кластеризации длинных текстов: {e}")
            logger.info("Пытаемся использовать метод для коротких текстов как fallback...")

            # Используем метод для коротких текстов
            return self.cluster_short_texts(id_text_dict, n_clusters, method, **kwargs)

    # ===== ОБЩИЕ МЕТОДЫ =====

    def _prepare_basic_data(self, id_text_dict: Dict[Any, str]) -> None:
        """Базовая подготовка данных"""
        self.document_ids = list(id_text_dict.keys())
        self.documents = list(id_text_dict.values())

        if len(self.documents) < 2:
            raise ValueError("Для кластеризации необходимо как минимум 2 документа")

    def _extract_keywords(self, text: str, top_k: int = 10) -> List[str]:
        """Извлекает ключевые слова из текста"""
        words = re.findall(r'\b[a-zA-Zа-яА-Я]{3,}\b', text.lower())
        self._extract_proper_nouns(text)
        words = [word for word in words if word not in self.stop_words and len(word) > 2]

        word_freq = Counter(words)
        return [word for word, freq in word_freq.most_common(top_k)]

    def _extract_proper_nouns(self, text):
        doc = self.nlp(text)
        for ent in doc.ents:
            if ent.label_ in ["PER", "ORG", "LOC", "GPE"]:  # Person, Organization, Location, Geo-Political
                self.stop_words.add(ent.text.lower())

    def _find_optimal_eps(self, embeddings: np.ndarray, min_samples: int = 2) -> float:
        """Автоматический подбор оптимального eps"""
        neighbors = NearestNeighbors(n_neighbors=min_samples)
        neighbors_fit = neighbors.fit(embeddings)
        distances, indices = neighbors_fit.kneighbors(embeddings)

        k_distances = np.sort(distances[:, min_samples - 1], axis=0)
        gradients = np.gradient(k_distances)
        second_derivatives = np.gradient(gradients)
        elbow_index = np.argmax(np.abs(second_derivatives))
        optimal_eps = k_distances[elbow_index]

        logger.info(f"Автоматически определенный eps: {optimal_eps:.3f}")
        return optimal_eps

    def _find_optimal_clusters(self, max_clusters: int) -> int:
        """Поиск оптимального количества кластеров"""
        if len(self.embeddings) < 2:
            return 1

        max_clusters = min(max_clusters, len(self.embeddings) - 1)
        best_score = -1
        best_n = 2

        for n in range(2, max_clusters + 1):
            kmeans = KMeans(n_clusters=n, random_state=42)
            labels = kmeans.fit_predict(self.embeddings)

            if len(set(labels)) > 1:
                score = silhouette_score(self.embeddings, labels)
                if score > best_score:
                    best_score = score
                    best_n = n

        logger.info(f"Оптимальное количество кластеров: {best_n} (silhouette score: {best_score:.3f})")
        return best_n

    def _format_results(self) -> Dict[str, Any]:
        """Форматирование результатов кластеризации"""
        if self.cluster_labels is None:
            raise ValueError("Сначала выполните кластеризацию")

        clusters = {}
        for doc_id, doc_text, cluster_id in zip(self.document_ids, self.documents, self.cluster_labels):
            if cluster_id not in clusters:
                clusters[cluster_id] = []

            keywords = self._extract_keywords(doc_text, top_k=5)

            clusters[cluster_id].append({
                'id': doc_id,
                'text': doc_text,
                'keywords': keywords,
                'text_length': len(doc_text),
                'processing_mode': self.processing_mode
            })

        metrics = {}
        if len(set(self.cluster_labels)) > 1:
            try:
                metrics['silhouette_score'] = silhouette_score(self.embeddings, self.cluster_labels)
            except:
                metrics['silhouette_score'] = None

        result = {
            'clusters': clusters,
            'cluster_labels': dict(zip(self.document_ids, self.cluster_labels)),
            'metrics': metrics,
            'total_clusters': len(set(self.cluster_labels)),
            'cluster_sizes': {k: len(v) for k, v in clusters.items()},
            'processing_mode': self.processing_mode
        }

        if self.processing_mode == 'long' and self.chunk_info:
            result['chunk_info'] = f"Создано {len(self.chunk_info)} чанков"

        return result

    def get_cluster_members(self, cluster_id: int) -> List[Dict]:
        """Получить все элементы определенного кластера"""
        results = self._format_results()
        return results['clusters'].get(cluster_id, [])

    def get_document_cluster(self, document_id: Any) -> int:
        """Получить кластер для конкретного документа"""
        if self.cluster_labels is None:
            raise ValueError("Сначала выполните кластеризацию")

        try:
            idx = self.document_ids.index(document_id)
            return self.cluster_labels[idx]
        except ValueError:
            raise ValueError(f"Документ с ID {document_id} не найден")

    def get_processing_mode(self) -> str:
        """Возвращает текущий режим обработки"""
        return self.processing_mode

    def save_to_word(self, results: Dict[str, Any], filename: str,
                     include_full_texts: bool = False,
                     max_text_length: int = 500) -> str:
        """
        Сохраняет результаты кластеризации в Word документ

        :param results: результаты кластеризации
        :param filename: имя файла для сохранения
        :param include_full_texts: включать ли полные тексты документов
        :param max_text_length: максимальная длина текста для отображения
        :return: путь к сохраненному файлу
        """
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise ImportError("Для работы этого метода установите python-docx: pip install python-docx")

        # Создаем документ
        doc = Document()

        # Заголовок
        title = doc.add_heading('Результаты кластеризации текстов', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Метод обработки
        mode_info = doc.add_paragraph()
        mode_info.add_run('Метод обработки: ').bold = True
        mode_info.add_run(f"{results.get('processing_mode', 'unknown')}")

        # Дата и время
        date_info = doc.add_paragraph()
        date_info.add_run('Дата анализа: ').bold = True
        date_info.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        # Общая статистика
        stats_heading = doc.add_heading('Общая статистика', level=1)

        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        stats_table.autofit = True

        # Заполняем таблицу статистики
        stats_data = [
            ('Всего документов', str(len(results['cluster_labels']))),
            ('Всего кластеров', str(results['total_clusters'])),
            ('Silhouette Score', f"{results['metrics'].get('silhouette_score', 'N/A'):.3f}"),
            ('Режим обработки', results.get('processing_mode', 'unknown'))
        ]

        for i, (key, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = key
            stats_table.cell(i, 1).text = value
            stats_table.cell(i, 0).paragraphs[0].runs[0].bold = True

        # Размеры кластеров
        sizes_heading = doc.add_heading('Размеры кластеров', level=1)

        sizes_table = doc.add_table(rows=len(results['cluster_sizes']) + 1, cols=2)
        sizes_table.style = 'Light Grid Accent 1'

        # Заголовки таблицы размеров
        sizes_table.cell(0, 0).text = 'ID кластера'
        sizes_table.cell(0, 1).text = 'Количество документов'
        sizes_table.cell(0, 0).paragraphs[0].runs[0].bold = True
        sizes_table.cell(0, 1).paragraphs[0].runs[0].bold = True

        # Заполняем данные о размерах кластеров
        for i, (cluster_id, size) in enumerate(results['cluster_sizes'].items(), 1):
            sizes_table.cell(i, 0).text = str(cluster_id)
            sizes_table.cell(i, 1).text = str(size)

        # Детальная информация по кластерам
        clusters_heading = doc.add_heading('Детальная информация по кластерам', level=1)

        for cluster_id, documents in results['clusters'].items():
            # Заголовок кластера
            cluster_heading = doc.add_heading(f'Кластер {cluster_id} ({len(documents)} документов)', level=2)

            # Ключевые слова кластера (объединяем ключевые слова всех документов)
            all_keywords = []
            for doc_info in documents:
                all_keywords.extend(doc_info.get('keywords', []))

            keyword_counter = Counter(all_keywords)
            top_keywords = [word for word, count in keyword_counter.most_common(10)]

            keywords_para = doc.add_paragraph()
            keywords_para.add_run('Ключевые слова: ').bold = True
            keywords_para.add_run(', '.join(top_keywords))

            # Таблица документов в кластере
            if documents:
                docs_table = doc.add_table(rows=len(documents) + 1, cols=3)
                docs_table.style = 'Light Grid Accent 1'

                # Заголовки таблицы
                headers = ['ID документа', 'Длина текста', 'Текст (фрагмент)']
                for col, header in enumerate(headers):
                    docs_table.cell(0, col).text = header
                    docs_table.cell(0, col).paragraphs[0].runs[0].bold = True

                # Данные документов
                for row, doc_info in enumerate(documents, 1):
                    docs_table.cell(row, 0).text = str(doc_info['id'])
                    docs_table.cell(row, 1).text = str(doc_info.get('text_length', 'N/A'))

                    # Обрезаем текст если нужно
                    text = doc_info['text']
                    if not include_full_texts and len(text) > max_text_length:
                        text = text[:max_text_length] + '...'
                    docs_table.cell(row, 2).text = text

            doc.add_paragraph()  # Пустая строка между кластерами

        # Шумовые кластеры (для DBSCAN)
        noise_cluster_id = -1
        if noise_cluster_id in results['clusters']:
            noise_heading = doc.add_heading('Шумовые документы (не кластеризованы)', level=2)
            noise_docs = results['clusters'][noise_cluster_id]

            noise_para = doc.add_paragraph()
            noise_para.add_run(f'Найдено шумовых документов: {len(noise_docs)}').bold = True

            if noise_docs:
                noise_table = doc.add_table(rows=len(noise_docs) + 1, cols=2)
                noise_table.style = 'Light Grid Accent 1'

                noise_table.cell(0, 0).text = 'ID документа'
                noise_table.cell(0, 1).text = 'Длина текста'
                noise_table.cell(0, 0).paragraphs[0].runs[0].bold = True
                noise_table.cell(0, 1).paragraphs[0].runs[0].bold = True

                for row, doc_info in enumerate(noise_docs, 1):
                    noise_table.cell(row, 0).text = str(doc_info['id'])
                    noise_table.cell(row, 1).text = str(doc_info.get('text_length', 'N/A'))

        # Информация о чанках (для длинных текстов)
        if 'chunk_info' in results:
            chunk_heading = doc.add_heading('Информация о разбиении на чанки', level=2)
            chunk_para = doc.add_paragraph()
            chunk_para.add_run(results['chunk_info'])

        # Сохраняем документ
        if not filename.endswith('.docx'):
            filename += '.docx'

        doc.save(filename)
        logger.info(f"Результаты сохранены в файл: {filename}")
        return filename

    def save_cluster_comparison(self, multiple_results: Dict[str, Dict[str, Any]],
                                filename: str) -> str:
        """
        Сохраняет сравнение нескольких кластеризаций в Word документ

        :param multiple_results: словарь {название_эксперимента: результаты}
        :param filename: имя файла для сохранения
        :return: путь к сохраненному файлу
        """
        doc = Document()

        # Заголовок
        title = doc.add_heading('Сравнение кластеризаций', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Тацалица сравнения
        comparison_heading = doc.add_heading('Сравнительная таблица', level=1)

        # Создаем таблицу сравнения
        num_results = len(multiple_results)
        comp_table = doc.add_table(rows=num_results + 1, cols=5)
        comp_table.style = 'Light Grid Accent 1'

        # Заголовки
        headers = ['Эксперимент', 'Метод', 'Кластеров', 'Silhouette', 'Размеры кластеров']
        for col, header in enumerate(headers):
            comp_table.cell(0, col).text = header
            comp_table.cell(0, col).paragraphs[0].runs[0].bold = True

        # Данные
        for row, (exp_name, results) in enumerate(multiple_results.items(), 1):
            comp_table.cell(row, 0).text = exp_name
            comp_table.cell(row, 1).text = results.get('processing_mode', 'unknown')
            comp_table.cell(row, 2).text = str(results['total_clusters'])

            silhouette = results['metrics'].get('silhouette_score', 'N/A')
            comp_table.cell(row, 3).text = f"{silhouette:.3f}" if isinstance(silhouette, float) else str(silhouette)

            sizes = ', '.join([f"{k}:{v}" for k, v in results['cluster_sizes'].items()])
            comp_table.cell(row, 4).text = sizes

        # Сохраняем документ
        if not filename.endswith('.docx'):
            filename += '.docx'

        doc.save(filename)
        logger.info(f"Сравнение сохранено в файл: {filename}")
        return filename